"""Extracción de texto plano desde distintas fuentes de conocimiento.

Cada parser recibe los bytes (o una URL) y devuelve texto limpio. Los errores
se propagan para que el servicio de ingesta marque la fuente como 'failed'.
"""

from __future__ import annotations

import io

import httpx

MAX_URL_BYTES = 5 * 1024 * 1024


def parse_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(p.strip() for p in pages if p.strip())


def parse_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    blocks: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    # Tablas: cada fila como línea "celda | celda"
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                blocks.append(" | ".join(cells))
    return "\n".join(blocks)


def parse_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    out: list[str] = []
    for ws in wb.worksheets:
        out.append(f"# Hoja: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                out.append(" | ".join(cells))
    wb.close()
    return "\n".join(out)


async def parse_url(url: str) -> str:
    async with httpx.AsyncClient(timeout=15, follow_redirects=True) as client:
        resp = await client.get(url, headers={"User-Agent": "WolfiaxBot/1.0"})
        resp.raise_for_status()
        content = resp.content[:MAX_URL_BYTES]

    ctype = resp.headers.get("content-type", "")
    if "application/pdf" in ctype:
        return parse_pdf(content)

    from bs4 import BeautifulSoup

    soup = BeautifulSoup(content, "html.parser")
    for tag in soup(["script", "style", "noscript", "header", "footer", "nav"]):
        tag.decompose()
    text = soup.get_text(separator="\n")
    lines = [ln.strip() for ln in text.splitlines()]
    return "\n".join(ln for ln in lines if ln)


def parse_bytes(source_type: str, data: bytes) -> str:
    if source_type == "pdf":
        return parse_pdf(data)
    if source_type == "docx":
        return parse_docx(data)
    if source_type == "xlsx":
        return parse_xlsx(data)
    # text | faq | policy | catalog exportado como texto
    return data.decode("utf-8", errors="replace")
